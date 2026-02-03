# -*- coding: utf-8 -*-
"""
文档解析工具集

提供文档读取、API规范提取、规范验证等功能。
支持 .docx/.txt/.md 格式的接口文档解析。

Note: This tool is loaded dynamically from the api_testing skill.
"""

import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from agentscope.tool import ToolResponse
from agentscope.message import TextBlock

# Add project root to path for common module access
PROJECT_ROOT = Path(__file__).parent.parent.parent.parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

# Try to use ToolConfig for workspace-aware path resolution
try:
    from agent.tool.base.config import ToolConfig
    USE_TOOL_CONFIG = True
except ImportError:
    USE_TOOL_CONFIG = False

# Fallback storage paths (used when ToolConfig is not available)
STORAGE_CHAT_DIR = (PROJECT_ROOT / "storage" / "chat").resolve()


def _get_storage_dir() -> Path:
    """Get the storage directory for uploaded files."""
    if USE_TOOL_CONFIG:
        try:
            config = ToolConfig.get()
            return config.workspace / "storage" / "chat"
        except RuntimeError:
            pass
    return STORAGE_CHAT_DIR


def read_document(file_path: str) -> ToolResponse:
    """
    Read and parse document content from uploaded files.

    Supports .docx, .txt, .md formats. Extracts plain text content
    for further processing by other tools.

    Args:
        file_path: Relative path to file (relative to storage/chat directory)
                   Example: "user123/conv456/api_spec.docx"

    Returns:
        ToolResponse containing document text content or error message
    """
    try:
        base_dir = _get_storage_dir()
        target_path = (base_dir / file_path).resolve()

        # 路径安全校验
        try:
            target_path.relative_to(base_dir)
        except ValueError:
            return ToolResponse(
                content=[TextBlock(
                    type="text",
                    text=json.dumps({
                        "status": "error",
                        "error_code": "ACCESS_DENIED",
                        "message": f"Access denied: path traversal detected"
                    }, ensure_ascii=False)
                )]
            )

        if not target_path.exists():
            return ToolResponse(
                content=[TextBlock(
                    type="text",
                    text=json.dumps({
                        "status": "error",
                        "error_code": "FILE_NOT_FOUND",
                        "message": f"File not found: {file_path}"
                    }, ensure_ascii=False)
                )]
            )

        suffix = target_path.suffix.lower()
        content = ""

        if suffix == ".docx":
            content = _parse_docx(target_path)
        elif suffix in [".txt", ".md"]:
            content = target_path.read_text(encoding="utf-8")
        elif suffix == ".pdf":
            return ToolResponse(
                content=[TextBlock(
                    type="text",
                    text=json.dumps({
                        "status": "error",
                        "error_code": "UNSUPPORTED_FORMAT",
                        "message": "PDF format not yet supported. Please convert to .docx or .txt"
                    }, ensure_ascii=False)
                )]
            )
        else:
            return ToolResponse(
                content=[TextBlock(
                    type="text",
                    text=json.dumps({
                        "status": "error",
                        "error_code": "UNSUPPORTED_FORMAT",
                        "message": f"Unsupported file format: {suffix}. Supported: .docx, .txt, .md"
                    }, ensure_ascii=False)
                )]
            )

        return ToolResponse(
            content=[TextBlock(
                type="text",
                text=json.dumps({
                    "status": "success",
                    "file_path": file_path,
                    "format": suffix,
                    "content": content,
                    "content_length": len(content)
                }, ensure_ascii=False)
            )]
        )

    except Exception as e:
        return ToolResponse(
            content=[TextBlock(
                type="text",
                text=json.dumps({
                    "status": "error",
                    "error_code": "PARSE_ERROR",
                    "message": f"Failed to parse document: {str(e)}"
                }, ensure_ascii=False)
            )]
        )


def _parse_docx(file_path: Path) -> str:
    """解析 .docx 文件内容"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx library required. Install via: pip install python-docx")

    doc = Document(str(file_path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def extract_api_spec(document_text: str) -> ToolResponse:
    """
    Extract API specification from document text.

    Analyzes document text to extract structured API information including
    endpoint, method, parameters, and expected responses.

    Args:
        document_text: Plain text content from document (output of read_document)

    Returns:
        ToolResponse containing extracted API specification in JSON format

    Example output:
        {
            "status": "success",
            "api_specs": [
                {
                    "endpoint": "/calculate_rate",
                    "method": "POST",
                    "description": "Calculate loan interest rate",
                    "request": {
                        "term": {"type": "string", "enum": ["12", "24", "36"]},
                        "repayment_method": {"type": "string"},
                        "use_coupon": {"type": "integer", "enum": [0, 20]}
                    },
                    "response": {
                        "rate": {"type": "number"}
                    },
                    "business_rules": []
                }
            ]
        }
    """
    try:
        specs = []

        # 提取基本API信息的正则模式
        patterns = {
            # URL/端点模式
            "endpoint": [
                r"(?:接口|API|URL|地址|路径)[：:\s]*([/\w\-_]+)",
                r"(?:endpoint|path)[：:\s]*([/\w\-_]+)",
                r"(\/[a-zA-Z_][\w\-_/]*)"
            ],
            # HTTP方法模式
            "method": [
                r"(?:方法|Method|请求方式)[：:\s]*(GET|POST|PUT|DELETE|PATCH)",
                r"\b(GET|POST|PUT|DELETE|PATCH)\b"
            ]
        }

        # 提取端点
        endpoint = None
        for pattern in patterns["endpoint"]:
            match = re.search(pattern, document_text, re.IGNORECASE)
            if match:
                endpoint = match.group(1)
                break

        # 提取方法
        method = "POST"  # 默认
        for pattern in patterns["method"]:
            match = re.search(pattern, document_text, re.IGNORECASE)
            if match:
                method = match.group(1).upper()
                break

        # 提取参数信息
        request_params = _extract_parameters(document_text)

        # 提取业务规则
        business_rules = _extract_business_rules(document_text)

        # 提取响应信息
        response_schema = _extract_response_schema(document_text)

        if endpoint:
            spec = {
                "endpoint": endpoint,
                "method": method,
                "description": _extract_description(document_text),
                "request": request_params,
                "response": response_schema,
                "business_rules": business_rules
            }
            specs.append(spec)

        if not specs:
            return ToolResponse(
                content=[TextBlock(
                    type="text",
                    text=json.dumps({
                        "status": "warning",
                        "message": "No API specification found in document. Please provide structured API documentation.",
                        "api_specs": [],
                        "hints": [
                            "Document should contain endpoint/URL information",
                            "HTTP method (GET/POST/PUT/DELETE) should be specified",
                            "Parameter definitions with types are helpful"
                        ]
                    }, ensure_ascii=False)
                )]
            )

        return ToolResponse(
            content=[TextBlock(
                type="text",
                text=json.dumps({
                    "status": "success",
                    "api_specs": specs,
                    "extracted_count": len(specs)
                }, ensure_ascii=False)
            )]
        )

    except Exception as e:
        return ToolResponse(
            content=[TextBlock(
                type="text",
                text=json.dumps({
                    "status": "error",
                    "error_code": "EXTRACTION_ERROR",
                    "message": f"Failed to extract API spec: {str(e)}"
                }, ensure_ascii=False)
            )]
        )


def _extract_parameters(text: str) -> Dict[str, Any]:
    """从文本中提取参数信息"""
    params = {}

    # 参数模式：参数名（类型）：描述 或 参数名: 枚举值
    param_patterns = [
        # 匹配表格形式：参数名 | 类型 | 描述
        r"(\w+)\s*\|\s*(string|int|integer|number|boolean|float)\s*\|",
        # 匹配文字描述：参数名（类型）
        r"(\w+)\s*[（(]\s*(string|int|integer|number|boolean|float)\s*[）)]",
        # 匹配冒号形式：参数名: 值1/值2/值3
        r"(\w+)\s*[：:]\s*([^,\n]+(?:/[^,\n]+)+)"
    ]

    for pattern in param_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            param_name = match[0]
            type_or_values = match[1]

            if "/" in type_or_values:
                # 枚举值
                enum_values = [v.strip() for v in type_or_values.split("/")]
                params[param_name] = {
                    "type": "string",
                    "enum": enum_values
                }
            else:
                # 类型
                type_map = {
                    "string": "string",
                    "int": "integer",
                    "integer": "integer",
                    "number": "number",
                    "float": "number",
                    "boolean": "boolean"
                }
                params[param_name] = {
                    "type": type_map.get(type_or_values.lower(), "string")
                }

    return params


def _extract_business_rules(text: str) -> List[str]:
    """从文本中提取业务规则"""
    rules = []

    # 业务规则关键词
    rule_keywords = [
        r"(?:规则|rule|约束|constraint)[：:\s]*(.+?)(?:\n|$)",
        r"(?:当|if|若).+(?:时|then).+",
        r"(?:兜底|默认|default|最低|最高).+",
        r"(?:必须|不能|应该|禁止).+"
    ]

    for pattern in rule_keywords:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            rule = match.strip() if isinstance(match, str) else match
            if rule and len(rule) > 5:  # 过滤太短的匹配
                rules.append(rule)

    return list(set(rules))[:10]  # 去重并限制数量


def _extract_response_schema(text: str) -> Dict[str, Any]:
    """从文本中提取响应结构"""
    response = {}

    # 响应字段模式
    response_patterns = [
        r"(?:响应|response|返回)[^{]*\{([^}]+)\}",
        r"(?:返回|输出)[：:\s]*(\w+)\s*[（(]([^）)]+)[）)]"
    ]

    for pattern in response_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            # 简单解析
            content = match.group(1)
            field_pattern = r'"?(\w+)"?\s*[：:]\s*"?([^",\n]+)"?'
            fields = re.findall(field_pattern, content)
            for field_name, field_value in fields:
                # 推断类型
                if field_value.replace(".", "").isdigit():
                    response[field_name] = {"type": "number"}
                else:
                    response[field_name] = {"type": "string"}

    return response


def _extract_description(text: str) -> str:
    """提取API描述"""
    desc_patterns = [
        r"(?:描述|description|说明|功能)[：:\s]*(.+?)(?:\n|$)",
        r"(?:接口|API)\s*[：:]\s*(.+?)(?:\n|$)"
    ]

    for pattern in desc_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()[:200]

    return ""


def validate_api_spec(api_spec: str) -> ToolResponse:
    """
    Validate completeness of API specification.

    Checks if the API specification contains all required fields
    and returns validation results with suggestions.

    Args:
        api_spec: JSON string of API specification (output of extract_api_spec)

    Returns:
        ToolResponse containing validation results

    Example output:
        {
            "status": "success",
            "valid": true,
            "errors": [],
            "warnings": ["Missing response schema"],
            "completeness_score": 85
        }
    """
    try:
        spec = json.loads(api_spec) if isinstance(api_spec, str) else api_spec

        errors = []
        warnings = []
        score = 100

        # 必填字段检查
        required_fields = ["endpoint", "method"]
        for field in required_fields:
            if field not in spec or not spec[field]:
                errors.append(f"Missing required field: {field}")
                score -= 25

        # 方法值检查
        valid_methods = ["GET", "POST", "PUT", "DELETE", "PATCH", "HEAD", "OPTIONS"]
        if spec.get("method") and spec["method"].upper() not in valid_methods:
            errors.append(f"Invalid HTTP method: {spec['method']}")
            score -= 10

        # 端点格式检查
        endpoint = spec.get("endpoint", "")
        if endpoint and not endpoint.startswith("/"):
            warnings.append("Endpoint should start with '/'")
            score -= 5

        # 请求参数检查
        request = spec.get("request", {})
        if not request:
            warnings.append("No request parameters defined")
            score -= 10
        else:
            # 检查参数类型定义
            for param_name, param_def in request.items():
                if not isinstance(param_def, dict) or "type" not in param_def:
                    warnings.append(f"Parameter '{param_name}' missing type definition")
                    score -= 5

        # 响应模式检查
        response = spec.get("response", {})
        if not response:
            warnings.append("No response schema defined")
            score -= 10

        # 业务规则检查
        rules = spec.get("business_rules", [])
        if not rules:
            warnings.append("No business rules defined - consider adding validation rules")

        valid = len(errors) == 0
        score = max(0, score)

        return ToolResponse(
            content=[TextBlock(
                type="text",
                text=json.dumps({
                    "status": "success",
                    "valid": valid,
                    "errors": errors,
                    "warnings": warnings,
                    "completeness_score": score,
                    "suggestion": "Specification is ready for test case generation" if valid else "Please fix errors before generating test cases"
                }, ensure_ascii=False)
            )]
        )

    except json.JSONDecodeError as e:
        return ToolResponse(
            content=[TextBlock(
                type="text",
                text=json.dumps({
                    "status": "error",
                    "error_code": "INVALID_JSON",
                    "message": f"Invalid JSON format: {str(e)}"
                }, ensure_ascii=False)
            )]
        )
    except Exception as e:
        return ToolResponse(
            content=[TextBlock(
                type="text",
                text=json.dumps({
                    "status": "error",
                    "error_code": "VALIDATION_ERROR",
                    "message": f"Validation failed: {str(e)}"
                }, ensure_ascii=False)
            )]
        )
